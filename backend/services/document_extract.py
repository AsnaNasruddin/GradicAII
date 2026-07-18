"""
Smart document text extraction with OCR fallback.

- Digital PDFs (typed, copyable text) -> fast local extraction via pdfplumber.
- Scanned/handwritten PDFs (no text layer) -> pages rendered to images via
  pypdfium2 and transcribed verbatim with GPT-4o vision OCR.
- Image files (jpg/png/webp/...) -> GPT-4o vision OCR directly.

Every caller gets back plain text, regardless of what kind of file was uploaded.
"""

import os
import io
import base64
import hashlib
import pdfplumber
from openai import OpenAI

OCR_MODEL = "gpt-4o"
OCR_SEED = 7                # fixed seed for reproducible transcripts
# Deliberately NOT under uploads/ — that directory is served publicly via
# StaticFiles with no auth, and these transcripts can contain full marking
# scheme / handwritten answer sheet text.
OCR_CACHE_DIR = os.path.join("private_cache", "ocr_cache")
MAX_OCR_PAGES = 30          # hard cap to bound API cost per document
PAGES_PER_OCR_CALL = 6      # pages sent per vision request
MIN_CHARS_PER_PAGE = 50     # below this average, the PDF is treated as scanned
RENDER_SCALE = 2.0          # ~150 DPI for A4, enough for handwriting
MAX_IMAGE_SIDE = 2000       # px, keeps request payloads reasonable

IMAGE_MIMES = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
    ".gif": "image/gif",
}

OCR_PROMPT = """You are a high-accuracy OCR transcription engine. Transcribe ALL text visible in the following images of a student's submitted document (pages {first_page} to {last_page}).

Rules:
- Transcribe the handwriting/print EXACTLY as written. Do NOT fix spelling or grammar, do NOT summarize, do NOT add or omit content.
- Keep the original reading order and structure. Preserve headings, question numbers/labels (e.g. "Q1", "Question 2", "Ans 3"), bullet points and numbering exactly as the student wrote them.
- Start each page's transcription with a marker line: === Page N ===
- If a word is unreadable, write [illegible]. If a page is blank, write [blank page].
- For diagrams, figures or tables, add a short note like [diagram: sketch of ...] and transcribe any text inside them.
- Output ONLY the transcription, nothing else."""


def _client() -> OpenAI:
    return OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))


def _ocr_cache_path(abs_path: str) -> str:
    """Cache key is the SHA-256 of the file bytes, so identical files (even
    re-uploaded under new names) always reuse the exact same transcript."""
    with open(abs_path, "rb") as f:
        digest = hashlib.sha256(f.read()).hexdigest()
    return os.path.join(OCR_CACHE_DIR, f"{digest}.txt")


def _ocr_cache_read(cache_path: str) -> str:
    try:
        with open(cache_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception:
        return ""


def _ocr_cache_write(cache_path: str, text: str) -> None:
    try:
        os.makedirs(OCR_CACHE_DIR, exist_ok=True)
        with open(cache_path, "w", encoding="utf-8") as f:
            f.write(text)
    except Exception:
        pass


def _render_pdf_pages_to_b64(abs_path: str) -> list:
    """Render each PDF page to a base64 JPEG using pypdfium2 (no external binaries)."""
    import pypdfium2 as pdfium
    from PIL import Image

    pages = []
    doc = pdfium.PdfDocument(abs_path)
    try:
        for i in range(min(len(doc), MAX_OCR_PAGES)):
            bitmap = doc[i].render(scale=RENDER_SCALE)
            img = bitmap.to_pil().convert("RGB")
            if max(img.size) > MAX_IMAGE_SIDE:
                img.thumbnail((MAX_IMAGE_SIDE, MAX_IMAGE_SIDE), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            pages.append(base64.b64encode(buf.getvalue()).decode())
    finally:
        doc.close()
    return pages


def _ocr_pages(b64_pages: list, mime: str = "image/jpeg") -> str:
    """Send page images to GPT-4o vision in batches and return the combined transcript."""
    client = _client()
    transcripts = []
    for start in range(0, len(b64_pages), PAGES_PER_OCR_CALL):
        batch = b64_pages[start:start + PAGES_PER_OCR_CALL]
        content = [{
            "type": "text",
            "text": OCR_PROMPT.format(first_page=start + 1, last_page=start + len(batch)),
        }]
        for offset, b64 in enumerate(batch):
            content.append({"type": "text", "text": f"Image below is page {start + offset + 1}:"})
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"},
            })
        response = client.chat.completions.create(
            model=OCR_MODEL,
            messages=[{"role": "user", "content": content}],
            max_tokens=4096,
            temperature=0,
            seed=OCR_SEED,
        )
        text = (response.choices[0].message.content or "").strip()
        if text:
            transcripts.append(text)
    return "\n\n".join(transcripts)


def extract_document_text(file_path: str) -> str:
    """
    Extract text from an uploaded document (PDF or image), OCR-ing scanned or
    handwritten content when needed. Always returns plain text ("" on failure).
    """
    if not file_path:
        return ""
    abs_path = file_path.lstrip("/")
    if not os.path.exists(abs_path):
        return ""

    ext = os.path.splitext(abs_path)[1].lower()

    if ext == ".pdf":
        embedded = ""
        page_count = 1
        try:
            with pdfplumber.open(abs_path) as pdf:
                page_count = max(len(pdf.pages), 1)
                embedded = "\n".join((p.extract_text() or "") for p in pdf.pages)
        except Exception:
            pass
        embedded = embedded.strip()

        # Enough real text per page -> digital PDF, no OCR needed
        if len(embedded) >= MIN_CHARS_PER_PAGE * page_count:
            return embedded

        # Scanned/handwritten PDF -> OCR each page with vision (cached by file hash)
        if not os.getenv("OPENAI_API_KEY"):
            return embedded
        try:
            cache_path = _ocr_cache_path(abs_path)
            cached = _ocr_cache_read(cache_path)
            if cached:
                return cached
            b64_pages = _render_pdf_pages_to_b64(abs_path)
            ocr_text = _ocr_pages(b64_pages).strip()
            if ocr_text:
                _ocr_cache_write(cache_path, ocr_text)
                return ocr_text
        except Exception:
            pass
        return embedded

    if ext == ".docx":
        try:
            import docx
            document = docx.Document(abs_path)
            return "\n".join(p.text for p in document.paragraphs).strip()
        except Exception:
            return ""

    if ext in IMAGE_MIMES:
        if not os.getenv("OPENAI_API_KEY"):
            return ""
        try:
            cache_path = _ocr_cache_path(abs_path)
            cached = _ocr_cache_read(cache_path)
            if cached:
                return cached
            with open(abs_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            ocr_text = _ocr_pages([b64], mime=IMAGE_MIMES[ext]).strip()
            if ocr_text:
                _ocr_cache_write(cache_path, ocr_text)
            return ocr_text
        except Exception:
            return ""

    # Fallback: treat anything else as plain text
    try:
        with open(abs_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception:
        return ""
